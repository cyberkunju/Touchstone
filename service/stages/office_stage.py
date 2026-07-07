"""Office + CSV stage — native cell values, formulas, formats (05 section 4).

openpyxl for XLSX (values + formulas + number formats + merged ranges),
python-docx for DOCX (runs + tables), csv.Sniffer for delimited text.

Native reads carry zero OCR uncertainty (N1 gold) — the numbers ARE the
document. Downstream closure laws still apply (a ledger that doesn't sum is
flagged), but transcription silents are structurally impossible here.
"""

from __future__ import annotations

import csv
import io
from dataclasses import dataclass, field


@dataclass
class NativeCell:
    ref: str                 # 'A1'
    value: str               # rendered value, always a string
    formula: str | None      # '=SUM(...)' when the cell is computed
    number_format: str | None


@dataclass
class NativeSheet:
    name: str
    cells: list[NativeCell] = field(default_factory=list)
    merged: list[str] = field(default_factory=list)   # 'A1:B2' ranges

    def cell_map(self) -> dict[str, str]:
        return {c.ref: c.value for c in self.cells}


def _render_value(v: object) -> str:
    """Canonical string for a cell value; floats keep repr precision minus
    trailing float noise (openpyxl returns real Python numbers)."""
    if v is None:
        return ""
    if isinstance(v, float):
        if v == int(v) and abs(v) < 1e15:
            return str(int(v))
        return repr(v)
    return str(v)


def extract_xlsx(data: bytes) -> list[NativeSheet]:
    """All sheets: values (data_only pass) + formulas (raw pass), merged ranges."""
    import openpyxl

    wb_values = openpyxl.load_workbook(io.BytesIO(data), data_only=True, read_only=False)
    wb_raw = openpyxl.load_workbook(io.BytesIO(data), data_only=False, read_only=False)

    sheets: list[NativeSheet] = []
    for name in wb_values.sheetnames:
        ws_v = wb_values[name]
        ws_r = wb_raw[name]
        sheet = NativeSheet(name=name,
                            merged=[str(r) for r in ws_v.merged_cells.ranges])
        for row in ws_v.iter_rows():
            for cell in row:
                if cell.value is None:
                    continue
                raw = ws_r[cell.coordinate].value
                formula = raw if isinstance(raw, str) and raw.startswith("=") else None
                sheet.cells.append(NativeCell(
                    ref=cell.coordinate,
                    value=_render_value(cell.value),
                    formula=formula,
                    number_format=cell.number_format,
                ))
        sheets.append(sheet)
    return sheets


@dataclass
class DocxContent:
    paragraphs: list[str]
    tables: list[list[list[str]]]     # table -> rows -> cells
    image_count: int                  # embedded images (recurse via vision)


def extract_docx(data: bytes) -> DocxContent:
    import docx

    document = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables = [
        [[cell.text for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]
    image_count = sum(
        1 for rel in document.part.rels.values()
        if "image" in rel.reltype
    )
    return DocxContent(paragraphs=paragraphs, tables=tables,
                       image_count=image_count)


def extract_csv(data: bytes) -> list[list[str]]:
    """Dialect-sniffed rows. Decoding tries UTF-8 (with BOM) then latin-1."""
    for enc in ("utf-8-sig", "latin-1"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:  # pragma: no cover — latin-1 never raises
        raise ValueError("undecodable text payload")

    sample = text[:8192]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
    except csv.Error:
        dialect = csv.excel
    return [row for row in csv.reader(io.StringIO(text), dialect) if row]
